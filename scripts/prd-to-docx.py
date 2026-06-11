# -*- coding: utf-8 -*-
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_el = tcPr.find(qn(f'w:{edge}'))
            if edge_el is None:
                edge_el = docx.oxml.OxmlElement(f'w:{edge}')
                tcPr.append(edge_el)
            edge_el.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            edge_el.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), kwargs[edge].get('color', '000000'))

def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.font.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10.5)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
    return table

def parse_md(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else '未命名'

    sections = {}
    current_sec = None
    for line in content.split('\n'):
        m = re.match(r'^##\s+\d+\.\s+(.+)$', line)
        if m:
            current_sec = m.group(1).strip()
            sections[current_sec] = []
            continue
        m = re.match(r'^###\s+\d+\.\d+\s+(.+)$', line)
        if m:
            current_sec = m.group(1).strip()
            sections[current_sec] = []
            continue
        if current_sec is not None:
            sections[current_sec].append(line)

    for sec in sections:
        sections[sec] = '\n'.join(sections[sec]).strip()

    return title, sections, content

def extract_table(text):
    rows = []
    lines = text.split('\n')
    in_table = False
    for line in lines:
        if line.strip().startswith('|'):
            parts = [p.strip() for p in line.split('|')[1:-1]]
            if not all(re.match(r'^[-:]+$', p) for p in parts):
                rows.append(parts)
    return rows

def extract_bullets(text):
    items = []
    for line in text.split('\n'):
        m = re.match(r'^-\s+(.+)$', line.strip())
        if m:
            items.append(m.group(1))
    return items

def build_docx(title, sections, raw_content, out_path):
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)

    # 封面标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.font.bold = True
    doc.add_paragraph()

    # 文档属性表
    add_heading(doc, '文档属性', level=2)
    add_table(doc,
        ['项目', '内容'],
        [
            ['文件状态', '[√] 发布'],
            ['当前版本', 'V1.0.0'],
            ['归属项目', '满改汽车改装平台'],
            ['编写人', ''],
            ['编写日期', '2026.06.10'],
            ['文档密级', '机密'],
        ]
    )
    doc.add_paragraph()

    # 版本历史
    add_heading(doc, '版本历史', level=2)
    add_table(doc,
        ['版本', '日期', '修订人', '修订说明'],
        [
            ['V1.0.0', '2026.06.10', '', '初稿发布'],
        ]
    )
    doc.add_page_break()

    # 一、引言
    add_heading(doc, '一、引言', level=1)

    add_heading(doc, '1、编写目的', level=2)
    overview = sections.get('页面概述', '')
    if overview:
        add_paragraph(doc, overview.replace('\n', ''), indent=True)
    else:
        add_paragraph(doc, f'本文编写的目的是提供一个开发系统的用户需求目标，并对所实现的 {title} 业务功能做全面的需求描述，作为系统设计和实现的目标及验收依据。')

    add_heading(doc, '2、背景', level=2)
    add_paragraph(doc, '满改平台是一个高端汽车改装服务系统，涵盖用户端 App、服务商端 App、平台网页端、品牌网页端和管理员端 App。')

    add_heading(doc, '3、定义', level=2)
    add_heading(doc, '3.1 术语', level=3)
    add_paragraph(doc, 'PRD：产品需求文档（Product Requirements Document）。')
    add_paragraph(doc, 'PV：页面浏览量（Page View）。')
    add_heading(doc, '3.2 缩略语', level=3)
    add_paragraph(doc, 'App：应用程序（Application）。')
    add_paragraph(doc, 'SKU：库存量单位（Stock Keeping Unit）。')

    add_heading(doc, '4、参考资料', level=2)
    add_paragraph(doc, '《满改平台App原型优化修改任务清单》')
    add_paragraph(doc, '《用户、服务商端App原型优化修改意见》')

    # 二、任务概述
    add_heading(doc, '二、任务概述', level=1)

    add_heading(doc, '1、产品定位', level=2)
    add_paragraph(doc, '满改平台是面向高端汽车改装领域的综合服务平台，连接车主用户、改装服务商和品牌供应商，提供从内容种草、方案咨询、商品采购到施工履约的一站式改装服务。')

    add_heading(doc, '2、产品目标', level=2)
    add_paragraph(doc, '通过平台化运营，提升改装行业的服务标准化水平和交易透明度，降低用户决策成本，帮助服务商拓展客源，实现多方共赢。')

    add_heading(doc, '3、目标用户与使用场景', level=2)
    add_paragraph(doc, '目标用户包括：有改装需求的车主、提供改装服务的服务商门店、管理平台的运营人员、品牌供应商。')

    add_heading(doc, '4、用户的特点', level=2)
    add_paragraph(doc, '车主用户注重改装品质与安全性，决策周期较长，需要案例参考和专业建议。服务商关注订单获取效率与结算周期。')

    add_heading(doc, '5、岗位与角色', level=2)
    roles = extract_table(sections.get('用户角色', ''))
    if len(roles) > 1:
        headers = roles[0]
        data = roles[1:]
        add_table(doc, headers, data)
    else:
        add_table(doc,
            ['角色/岗位名称', '职责', '权限', '描述'],
            [
                ['普通用户', '浏览与交易', '全部功能', '已登录用户'],
                ['游客', '浏览', '部分功能受限', '未登录用户'],
            ]
        )

    add_heading(doc, '6、假定和约束', level=2)
    add_paragraph(doc, '当前实现为前端静态原型，所有数据仅保存在浏览器内存或 localStorage 中，刷新页面后恢复初始 mock 数据。')

    # 三、产品结构
    add_heading(doc, '三、产品结构', level=1)
    add_paragraph(doc, '（按需绘制产品架构图、功能结构图、信息结构图）')

    # 四、全局说明
    add_heading(doc, '四、全局说明', level=1)
    add_heading(doc, '1、全局交互规范', level=2)
    add_paragraph(doc, '列表加载中/列表没有更多数据/消息没有更多数据：页面底部展示对应提示文案。')
    add_paragraph(doc, '弱网加载/请求或加载超时：提示网络异常，支持重试。')
    add_paragraph(doc, '列表没有数据：展示空状态插图与提示文案。')

    # 五、功能需求
    add_heading(doc, '五、功能需求', level=1)

    # 功能描述
    add_heading(doc, '1、功能描述', level=2)
    desc = sections.get('功能描述', '')
    if desc:
        add_paragraph(doc, desc, indent=True)
    else:
        add_paragraph(doc, f'{title} 的核心功能概述。')

    # 功能分析
    add_heading(doc, '2、功能分析', level=2)
    bullets = extract_bullets(sections.get('功能分析', ''))
    for b in bullets:
        add_paragraph(doc, f'· {b}', indent=True)
    if not bullets:
        add_paragraph(doc, '（功能点分析待补充）', indent=True)

    # 用户角色
    add_heading(doc, '3、用户角色', level=2)
    role_text = sections.get('用户角色', '')
    if role_text:
        roles = extract_table(role_text)
        if len(roles) > 1:
            add_table(doc, roles[0], roles[1:])
        else:
            add_paragraph(doc, role_text, indent=True)
    else:
        add_table(doc,
            ['角色/岗位名称', '职责', '权限', '描述'],
            [['普通用户', '浏览与交易', '全部功能', '已登录用户']]
        )

    # 数据字典
    add_heading(doc, '4、数据字典', level=2)
    ds_text = sections.get('数据源', '')
    if ds_text:
        ds = extract_table(ds_text)
        if len(ds) > 1:
            add_table(doc, ds[0], ds[1:])
        else:
            add_paragraph(doc, '（数据字典待补充）', indent=True)
    else:
        add_paragraph(doc, '（数据字典待补充）', indent=True)

    # 功能详细说明
    add_heading(doc, '5、功能详细说明', level=2)

    # 从 raw_content 中提取所有 6.x 子章节
    detail_sections = re.findall(r'###\s+\d+\.\d+\s+(.+?)\n(.*?)(?=###\s+\d+\.\d+|##\s+\d+\.|\Z)', raw_content, re.DOTALL)
    if not detail_sections:
        detail_sections = re.findall(r'###\s+(.+?)\n(.*?)(?=###\s+|##\s+\d+\.|\Z)', raw_content, re.DOTALL)

    for idx, (sub_title, sub_content) in enumerate(detail_sections, 1):
        add_heading(doc, f'5.{idx} {sub_title.strip()}', level=3)

        # 提取用户场景、前置条件、需求描述、后置条件、补充说明
        for label in ['用户场景', '前置条件', '需求描述', '后置条件', '补充说明']:
            m = re.search(rf'\*\*{label}\*\*\n(.*?)(?=\*\*|\Z)', sub_content, re.DOTALL)
            if m:
                text = m.group(1).strip().replace('\n', ' ')
                p = doc.add_paragraph()
                run = p.add_run(f'{label}：')
                run.font.bold = True
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10.5)
                run2 = p.add_run(text)
                run2.font.name = '宋体'
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run2.font.size = Pt(10.5)
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.line_spacing = Pt(20)
            else:
                # 尝试匹配没有 ** 包裹的格式
                m = re.search(rf'{label}\n(.*?)(?=用户场景|前置条件|需求描述|后置条件|补充说明|\Z)', sub_content, re.DOTALL)
                if m:
                    text = m.group(1).strip().replace('\n', ' ')
                    add_paragraph(doc, f'{label}：{text}', indent=True)

    if not detail_sections:
        add_paragraph(doc, '（功能详细说明待补充）', indent=True)

    doc.save(out_path)
    print(f'Created: {out_path}')

def main():
    prd_dir = os.path.join(os.path.dirname(__file__), '..', 'docs', 'prd')
    out_dir = os.path.join(os.path.dirname(__file__), '..', 'docs', 'prd-docx-v2')
    os.makedirs(out_dir, exist_ok=True)

    for entry in os.listdir(prd_dir):
        sub_dir = os.path.join(prd_dir, entry)
        if not os.path.isdir(sub_dir):
            continue
        out_sub = os.path.join(out_dir, entry)
        os.makedirs(out_sub, exist_ok=True)
        for file in os.listdir(sub_dir):
            if not file.endswith('.md'):
                continue
            md_path = os.path.join(sub_dir, file)
            out_path = os.path.join(out_sub, file.replace('.md', '.docx'))
            title, sections, raw = parse_md(md_path)
            build_docx(title, sections, raw, out_path)

    print('\nAll done.')

if __name__ == '__main__':
    main()
